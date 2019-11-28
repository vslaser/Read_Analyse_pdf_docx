import PyPDF2
import pandas as pd 
import numpy as np 
import os
import glob

def readpdftext(filepath):
    print("reading pdf as text file....")
    pdfFileObj = open(filepath,'rb')     #'rb' for read binary mode
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    pdfReader.numPages
    pdffile = '' # This creates an empty string object to which the text from the pdf file will be added
    for i in range(0,pdfReader.numPages): # iterating page by page to read text
        pageObj = pdfReader.getPage(i)          
        pdffile = pdffile + pageObj.extractText()
    print("file ready. file format: ", type(pdffile))
    return pdffile # when the function is called, it returns text of the pdf file and assigns to the variable used


def readpdffiles(filespath):
    pdffiles =  glob.glob(os.path.join(filespath, '*.pdf')) # identify the list of all pdf documents
    print("reading %x pdf files from the given path...", len(pdffiles))
    data = {} # creating a dictionary which will have the file names as keys and content as value
    # Dictionaries are key:value pairs and the value can be accessed using the key 
    for file in pdffiles: 
        data.update({os.path.splitext(file)[0]: readpdftext(file)})
    print("All pdf files read into a dictionary with file path and text")
    return data

def writepdfs_to_docx(data):
    keys = data.keys()
    data2 = ''
    for k in keys:
        text = str(k) + str(data[k].encode('utf-8'))
        data2 = data2 + text
    from docx import Document
    from docx.shared import Inches
    document = Document()
    document.add_paragraph(data2)
    document.save('Document.docx')
    print("Pdf files written to docx")

def singlepdf_to_docx(data):
    from docx import Document
    from docx.shared import Inches
    document = Document()
    document.add_paragraph(data)
    document.save('Document.docx')
    print("Pdf file written to docx")

'''def table_from_pdf_page(filepath):
    print("Trying to read all tables from the pdf file....")
    import tabula
    df = []
    with open(filepath, mode = "rb") as f:
        reader = PyPDF2.PdfFileReader(f)
        n = reader.getNumPages()
    for i in range(0,n+1):
        df.append(tabula.read_pdf(filepath, multiple_tables = True, pages = i))
    data = []
    pages = []
    for i in range(0,len(df)):
        if len(df[i]) != 0:
            data.append(df[i])
            pages.append(i)

    print("Number of tables read in pdf file: ", len(data))
    print("Number of pages in which tables were read: ", len(pages))
    return data, pages'''

def tables_from_singlepdf(filepath):
    import camelot # make sure to install camelot-py[cv] via pip also the system needs ghostscript correctly installed

    tables = camelot.read_pdf(filepath, pages='1,2-end')
    print("Number of Tables: ", len(tables))
    return tables

def wordcloud_single_pdf(filepath):
    data = readpdftext(filepath)
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    from nltk.tokenize import RegexpTokenizer
    from nltk.tokenize import sent_tokenize
    import string
    from nltk.stem.porter import PorterStemmer
    
    stop_words = set(stopwords.words("english"))

    data = data.translate(str.maketrans('', '', string.punctuation))
    data = data.replace("\n", " ")
    data = data.lower()
    tokenized_sent = sent_tokenize(str(data))
    tokenizer = RegexpTokenizer('\n', gaps=True)
    tokenized_word = tokenizer.tokenize(str(tokenized_sent))
    stemmer = PorterStemmer()
    tokens_stemmed = [stemmer.stem(x) for x in tokenized_word]


    filtered_words = []
    for w in tokens_stemmed:
        if w not in stop_words:
            filtered_words.append(w)
        
    from wordcloud import WordCloud
    import matplotlib.pyplot as plt

    wordcloud = WordCloud(width = 1000, height = 500).generate(" ".join(filtered_words))
    plt.figure(figsize = (15,8))
    plt.imshow(wordcloud)
    plt.axis("off")
    head, tail = os.path.split(filepath)
    name = tail.replace(".pdf", "") + ".jpg"
    plt.savefig(name)

    print("Word Cloud Saved")

def word_freq(filespath):
    def pdftext(filepath):
    
        pdfFileObj = open(filepath,'rb')     #'rb' for read binary mode
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        pdfReader.numPages
        pdffile = '' # This creates an empty string object to which the text from the pdf file will be added
        for i in range(0,pdfReader.numPages): # iterating page by page to read text
            pageObj = pdfReader.getPage(i)          
            pdffile = pdffile + pageObj.extractText()
        
        return pdffile # when the function is called, it returns text of the pdf file and assigns to the variable used


    def pdffiles(filespath):
        pdffiles =  glob.glob(os.path.join(filespath, '*.pdf')) # identify the list of all pdf documents
        
        data = {} # creating a dictionary which will have the file names as keys and content as value
        # Dictionaries are key:value pairs and the value can be accessed using the key 
        for file in pdffiles: 
            data.update({os.path.splitext(file)[0]: pdftext(file)})
        
        return data

    d = pdffiles(filespath)
    n = input("Number of words: ")
    cols = ["PDF"]
    strng = ""
    for i in range(int(n)):
        x = input()
        strng = strng+ x + " "
        cols.append(x)
    strng = strng[:-1]
    
    cols.append(strng)
    dat2 = pd.DataFrame(columns = cols)
    for k in range(len(d.keys())):
        d[list(d.keys())[k]] = d[list(d.keys())[k]].lower()
        words = d[list(d.keys())[k]].split(" ")
        
        head , tail = os.path.split(list(d.keys())[k])
        lst = [tail]
        for i in range(1, len(cols)):
            lst.append(words.count(cols[i].lower()))
            
        temp = pd.DataFrame([lst])
        temp.columns = cols
        dat2 = pd.concat([dat2,temp], axis = 0, ignore_index = True)
    return dat2


def word_freq_one_pdf(filespath):
    def pdftext(filepath):
    
        pdfFileObj = open(filepath,'rb')     #'rb' for read binary mode
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        pdfReader.numPages
        pdffile = '' # This creates an empty string object to which the text from the pdf file will be added
        for i in range(0,pdfReader.numPages): # iterating page by page to read text
            pageObj = pdfReader.getPage(i)          
            pdffile = pdffile + pageObj.extractText()
        
        return pdffile # when the function is called, it returns text of the pdf file and assigns to the variable used


    d = pdftext(filespath)
    n = input("Number of words: ")
    cols = ["PDF"]
    strng = ""
    for i in range(int(n)):
        x = input()
        strng = strng+ x + " "
        cols.append(x)
    strng = strng[:-1]
    
    cols.append(strng)
    dat2 = pd.DataFrame(columns = cols)
    d = d.lower()
    words = d.split(" ")
        
    head , tail = os.path.split(filespath)
    lst = [tail]
    for i in range(1, len(cols)):
        lst.append(words.count(cols[i].lower()))
        
    temp = pd.DataFrame([lst])
    temp.columns = cols
    dat2 = pd.concat([dat2,temp], axis = 0, ignore_index = True)
    return dat2


def ngram_freq(pdffilepath):
    import nltk
    from nltk.util import ngrams
    import collections
    import string
    from nltk.corpus import stopwords 
    import re

    txt = readpdftext(pdffilepath)
    txt = txt.lower()
    txt = txt.replace("\n", " ")
    txt = re.sub('<.*>','',txt)
    txt = txt.translate(str.maketrans('', '', string.punctuation))
    tokens = nltk.word_tokenize(txt)
    stops = set(stopwords.words('english'))
    filtered = [w for w in tokens if not w in stops]

    n=input("ngram number: ")
    n = int(n)
    bgs = ngrams(filtered,n)
    bgsfreq = collections.Counter(bgs)
    #bgsfreq.most_common(100)[0][1]
    freqs = pd.DataFrame(columns = ["Words", "Frequency"])
    for i in range(len(bgsfreq.most_common(100))):
        lst = []
        lst.append(bgsfreq.most_common(100)[i][0])
        lst.append(bgsfreq.most_common(100)[i][1])
        temp = pd.DataFrame([lst])
        temp.columns = ["Words", "Frequency"]
        freqs = pd.concat([freqs,temp], axis = 0)
    return freqs
    



